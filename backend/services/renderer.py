import io, os, re
from reportlab.lib.pagesizes import A5
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

THEMES = {
    "blue":   {"primary":"#2563EB","light":"#EFF6FF","accent":"#1D4ED8","text":"#1E3A5F"},
    "green":  {"primary":"#16A34A","light":"#F0FDF4","accent":"#15803D","text":"#14532D"},
    "orange": {"primary":"#EA580C","light":"#FFF7ED","accent":"#C2410C","text":"#7C2D12"},
    "purple": {"primary":"#7C3AED","light":"#F5F3FF","accent":"#6D28D9","text":"#3B0764"},
    "red":    {"primary":"#DC2626","light":"#FEF2F2","accent":"#B91C1C","text":"#7F1D1D"},
}
_FONT_CACHE = {}


def _hex(h):
    h = h.lstrip("#")
    return colors.Color(int(h[0:2],16)/255, int(h[2:4],16)/255, int(h[4:6],16)/255)

def _rgb(h):
    h = h.lstrip("#")
    return int(h[0:2],16), int(h[2:4],16), int(h[4:6],16)

def _clean(t):
    return re.sub(
        r"[\U0001F300-\U0001F9FF\U00002600-\U000027BF\U0001FA00-\U0001FAFF]+",
        "", str(t), flags=re.UNICODE
    ).strip()

def _get_fonts():
    if _FONT_CACHE:
        return _FONT_CACHE["fn"], _FONT_CACHE["fb"]
    candidates = [
        ("C:/Windows/Fonts/malgun.ttf",   "Malgun"),
        ("C:/Windows/Fonts/NanumGothic.ttf", "Nanum"),
    ]
    fn = "Helvetica"
    for path, name in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                fn = name
                break
            except Exception:
                pass
    _FONT_CACHE["fn"] = fn
    _FONT_CACHE["fb"] = fn   # bold fallback = same font
    return fn, fn


def _wrap(c, text, x, y, w, fn, sz, col, lh=None, max_lines=None):
    """텍스트 자동 줄바꿈. max_lines 지정 시 넘치면 '...' 처리."""
    if not text:
        return y
    lh = lh or sz * 1.45
    c.setFont(fn, sz)
    c.setFillColor(col)
    words = text.split()
    line = ""
    drawn = 0
    for word in words:
        test = (line + " " + word).strip()
        if c.stringWidth(test, fn, sz) <= w:
            line = test
        else:
            if line:
                if max_lines and drawn >= max_lines - 1:
                    # 마지막 줄: ... 추가 후 종료
                    ellipsis_line = line
                    while ellipsis_line and c.stringWidth(ellipsis_line + "...", fn, sz) > w:
                        ellipsis_line = ellipsis_line[:-1]
                    c.drawString(x, y, ellipsis_line + "...")
                    y -= lh
                    return y
                c.drawString(x, y, line)
                y -= lh
                drawn += 1
            line = word
    if line:
        c.drawString(x, y, line)
        y -= lh
    return y


def _estimate_wrap_height(c, text, w, fn, sz, lh=None):
    """텍스트 블록 예상 높이(pt) 계산"""
    if not text:
        return 0
    lh = lh or sz * 1.45
    c.setFont(fn, sz)
    words = text.split()
    line = ""
    lines = 0
    for word in words:
        test = (line + " " + word).strip()
        if c.stringWidth(test, fn, sz) <= w:
            line = test
        else:
            lines += 1
            line = word
    if line:
        lines += 1
    return lines * lh


# ─── PDF 렌더링 ─────────────────────────────────────────────

def render_pdf(data: dict) -> bytes:
    buf = io.BytesIO()
    fn, fb = _get_fonts()

    th  = THEMES.get(data.get("color_theme", "blue"), THEMES["blue"])
    pr  = _hex(th["primary"])
    li  = _hex(th["light"])
    ac  = _hex(th["accent"])
    tx  = _hex(th["text"])
    W, H = A5                   # 419.53 × 595.28 pt
    lm   = 13 * mm
    rm   = W - 13 * mm
    iw   = rm - lm              # inner width
    bot  = 12 * mm              # 하단 여백 한계

    c = pdf_canvas.Canvas(buf, pagesize=A5)
    page_num = [1]
    y = H   # 초기화 (실제 사용 전 덮어쓰임)

    # ── 공통: 2페이지 이후 소형 헤더
    def mini_header():
        c.setFillColor(pr)
        c.rect(0, H - 11*mm, W, 11*mm, fill=1, stroke=0)
        c.setFillColor(ac)
        c.rect(0, H - 11*mm, 2.5, 11*mm, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont(fb, 8)
        c.drawString(lm + 2*mm, H - 7.5*mm, _clean(data.get("title",""))[:34])
        c.setFont(fn, 6)
        c.setFillColor(colors.Color(1,1,1,0.55))
        c.drawRightString(rm, H - 7.5*mm, f"p.{page_num[0]}  ·  A5 REPORT")

    def new_page():
        nonlocal y
        c.showPage()
        page_num[0] += 1
        mini_header()
        y = H - 11*mm - 7*mm

    def check(needed):
        if y - needed < bot:
            new_page()

    # ═══════════════════════════════════════
    #  PAGE 1 : 제목 헤더 + 요약 + 목차 미리보기
    # ═══════════════════════════════════════
    hh = 40 * mm
    c.setFillColor(pr)
    c.rect(0, H - hh, W, hh, fill=1, stroke=0)
    c.setFillColor(ac)
    c.rect(0, H - hh, 3*mm, hh, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont(fb, 16)
    c.drawString(lm + 2*mm, H - 17*mm, _clean(data.get("title",""))[:26])
    sub = _clean(data.get("subtitle",""))
    if sub:
        c.setFont(fn, 8.5)
        c.setFillColor(colors.Color(1,1,1,0.82))
        c.drawString(lm + 2*mm, H - 26*mm, sub[:55])
    c.setFont(fn, 6)
    c.setFillColor(colors.Color(1,1,1,0.5))
    c.drawRightString(rm, H - hh + 3.5*mm, "A5 REPORT  |  Claude AI")
    y = H - hh - 5*mm

    # ── 요약 박스 ──────────────────────────────
    sm = _clean(data.get("summary",""))
    if sm:
        sm_h = max(22*mm, _estimate_wrap_height(c, sm, iw - 8*mm, fn, 8.5, 12) + 14*mm)
        check(sm_h)
        c.setFillColor(li)
        c.roundRect(lm, y - sm_h, iw, sm_h, 3, fill=1, stroke=0)
        c.setFillColor(pr)
        c.rect(lm, y - sm_h, 2.5, sm_h, fill=1, stroke=0)
        c.setFont(fb, 6.5); c.setFillColor(ac)
        c.drawString(lm + 5*mm, y - 5*mm, "SUMMARY")
        _wrap(c, sm, lm + 5*mm, y - 11*mm, iw - 8*mm, fn, 8.5, tx, 12)
        y -= sm_h + 6*mm

    # ── 목차 (제목 + 내용 미리보기) ────────────
    kps = data.get("key_points", [])
    if kps:
        check(10*mm)
        c.setFont(fb, 7); c.setFillColor(ac)
        c.drawString(lm, y, "CONTENTS")
        y -= 5*mm

        for i, kp in enumerate(kps, 1):
            hd  = _clean(kp.get("heading",""))
            ct  = _clean(kp.get("content",""))

            # 미리보기: 3줄까지만
            preview_w  = iw - 10*mm
            preview_h  = min(
                _estimate_wrap_height(c, ct, preview_w, fn, 8, 11),
                3 * 8 * 1.45          # 최대 3줄
            )
            item_h = 7*mm + preview_h + 3*mm
            check(item_h)

            # 번호 원
            c.setFillColor(pr)
            c.circle(lm + 2.5*mm, y - 2.5*mm, 2.5*mm, fill=1, stroke=0)
            c.setFont(fb, 7.5); c.setFillColor(colors.white)
            c.drawCentredString(lm + 2.5*mm, y - 4*mm, str(i))

            # 소제목
            c.setFont(fb, 9.5); c.setFillColor(ac)
            c.drawString(lm + 7*mm, y - 3.5*mm, hd[:36])
            y -= 7*mm

            # 내용 미리보기 (3줄 제한)
            _wrap(c, ct, lm + 7*mm, y, preview_w, fn, 8,
                  colors.Color(0.35,0.35,0.35), 11, max_lines=3)
            y -= preview_h + 3*mm

    # ═══════════════════════════════════════
    #  PAGE 2+ : 상세 본문
    # ═══════════════════════════════════════
    c.showPage()
    page_num[0] += 1
    mini_header()
    y = H - 11*mm - 7*mm

    c.setFont(fb, 7); c.setFillColor(ac)
    c.drawString(lm, y, "DETAILED CONTENT")
    y -= 6*mm

    for i, kp in enumerate(kps, 1):
        hd = _clean(kp.get("heading",""))
        ct = _clean(kp.get("content",""))
        ct_h = _estimate_wrap_height(c, ct, iw - 4*mm, fn, 8.5, 12)
        needed = 6*mm + 5.5*mm + ct_h + 8*mm   # 번호줄 + 소제목 + 본문 + 여백
        check(needed)

        # 번호 + 구분선
        num = f"{i:02d}"
        c.setFont(fb, 11); c.setFillColor(pr)
        nw = c.stringWidth(num, fb, 11)
        c.drawString(lm, y, num)
        c.setStrokeColor(colors.Color(0.80,0.80,0.80))
        c.setLineWidth(0.4)
        c.line(lm + nw + 2*mm, y - 1, rm, y - 1)
        y -= 6*mm

        # 소제목
        c.setFont(fb, 9.5); c.setFillColor(ac)
        c.drawString(lm + 2*mm, y, hd[:36])
        y -= 5.5*mm

        # 본문 (전문)
        y = _wrap(c, ct, lm + 2*mm, y, iw - 4*mm, fn, 8.5,
                  colors.Color(0.15,0.15,0.15), 12)
        y -= 8*mm

    # ═══════════════════════════════════════
    #  마지막 : 시각 인포그래픽
    # ═══════════════════════════════════════
    # 인포그래픽 최소 예상 높이 계산
    stats      = data.get("stats", [])
    stats_h    = (18*mm + 5*mm) if stats else 0
    takeaway_h = 15*mm + 4*mm if data.get("takeaway") else 0
    min_infographic = 14*mm + stats_h + 10*mm + takeaway_h  # 배너 + stats + 카드 첫행 + takeaway

    if y - min_infographic < bot:
        c.showPage()
        page_num[0] += 1
        mini_header()
        y = H - 11*mm - 7*mm
    else:
        y -= 5*mm
        c.setStrokeColor(colors.Color(0.85,0.85,0.85))
        c.setLineWidth(0.4)
        c.line(lm, y, rm, y)
        y -= 5*mm

    # VISUAL SUMMARY 배너
    c.setFillColor(pr)
    c.roundRect(lm, y - 10*mm, iw, 10*mm, 3, fill=1, stroke=0)
    c.setFont(fb, 8); c.setFillColor(colors.white)
    c.drawCentredString(W/2, y - 7*mm, "VISUAL SUMMARY")
    y -= 14*mm

    # 통계 박스 (있을 때만)
    if stats:
        sbh = 18*mm
        bw  = (iw - (len(stats)-1)*3*mm) / len(stats)
        xs  = lm
        check(sbh + 5*mm)
        for s in stats:
            c.setFillColor(li)
            c.roundRect(xs, y - sbh, bw, sbh, 4, fill=1, stroke=0)
            c.setFillColor(pr)
            c.roundRect(xs, y - sbh, bw, 4, 0, fill=1, stroke=0)
            c.setFont(fb, 14); c.setFillColor(pr)
            c.drawCentredString(xs + bw/2, y - 11*mm, _clean(str(s.get("value","")))[:10])
            c.setFont(fn, 6.5); c.setFillColor(tx)
            c.drawCentredString(xs + bw/2, y - 16*mm, _clean(str(s.get("label","")))[:20])
            xs += bw + 3*mm
        y -= sbh + 5*mm

    # 2열 카드 그리드 (행별 동적 높이)
    if kps:
        col_w = (iw - 3*mm) / 2
        card_top_pad  = 13 * mm   # 번호원 + 제목 영역
        card_bot_pad  = 3  * mm
        card_side_pad = 2  * mm
        text_w = col_w - card_side_pad * 2

        # 행별로 처리
        for row in range(0, len(kps), 2):
            left_kp  = kps[row]
            right_kp = kps[row+1] if row+1 < len(kps) else None

            # 각 카드 텍스트 예상 높이 계산 (제한 없이 전체 내용)
            lct = _clean(left_kp.get("content",""))
            lh_text = _estimate_wrap_height(c, lct, text_w, fn, 7.5, 10.5)
            rh_text = 0
            if right_kp:
                rct = _clean(right_kp.get("content",""))
                rh_text = _estimate_wrap_height(c, rct, text_w, fn, 7.5, 10.5)

            card_h = max(24*mm,
                         card_top_pad + lh_text + card_bot_pad,
                         card_top_pad + rh_text + card_bot_pad)

            check(card_h + 3*mm)

            for col_idx, kp in enumerate([left_kp] + ([right_kp] if right_kp else [])):
                col_x = lm + col_idx * (col_w + 3*mm)
                i_card = row + col_idx
                ct_card = _clean(kp.get("content",""))
                hd_card = _clean(kp.get("heading",""))

                # 카드 배경
                bg = li if col_idx == 0 else _hex(th["light"])
                # 두 카드 색상 약간 구분
                bg_colors = [
                    _hex(th["light"]),
                    colors.Color(
                        _hex(th["light"]).red   * 0.93,
                        _hex(th["light"]).green * 0.93,
                        _hex(th["light"]).blue  * 0.93,
                    )
                ]
                c.setFillColor(bg_colors[col_idx % 2])
                c.roundRect(col_x, y - card_h, col_w, card_h, 4, fill=1, stroke=0)

                # 상단 강조 바
                c.setFillColor(pr)
                c.roundRect(col_x, y - 3, col_w, 3, 0, fill=1, stroke=0)

                # 번호 원
                c.setFillColor(pr)
                c.circle(col_x + 4.5*mm, y - 6.5*mm, 3.5*mm, fill=1, stroke=0)
                c.setFont(fb, 8); c.setFillColor(colors.white)
                c.drawCentredString(col_x + 4.5*mm, y - 8.3*mm, str(i_card+1))

                # 소제목 (긴 제목도 처리)
                c.setFont(fb, 7.5); c.setFillColor(ac)
                c.drawString(col_x + 9.5*mm, y - 7.5*mm, hd_card[:22])

                # 본문 (전체 내용, 줄 제한 없음)
                _wrap(c, ct_card,
                      col_x + card_side_pad, y - card_top_pad,
                      text_w, fn, 7.5,
                      colors.Color(0.2,0.2,0.2), 10.5)

            y -= card_h + 3*mm

    # 결론 배너
    tk = _clean(data.get("takeaway",""))
    if tk:
        tk_h = max(15*mm, _estimate_wrap_height(c, tk, iw - 12*mm, fb, 9, 12) + 10*mm)
        check(tk_h + 3*mm)
        y -= 2*mm
        c.setFillColor(pr)
        c.roundRect(lm, y - tk_h, iw, tk_h, 3, fill=1, stroke=0)
        c.setFont(fb, 9); c.setFillColor(colors.white)
        _wrap(c, tk, lm + 5*mm, y - 5.5*mm, iw - 10*mm, fb, 9, colors.white, 12)

    c.save()
    return buf.getvalue()


# ─── DOCX 렌더링 ────────────────────────────────────────────

def render_docx(data: dict) -> bytes:
    th  = THEMES.get(data.get("color_theme","blue"), THEMES["blue"])
    pr  = _rgb(th["primary"])
    ac  = _rgb(th["accent"])
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(14.8)
    sec.page_height   = Cm(21.0)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(1.5)

    def ar(p, t, bold=False, sz=10, rgb=None):
        r = p.add_run(t)
        r.bold = bold
        r.font.size = Pt(sz)
        if rgb:
            r.font.color.rgb = RGBColor(*rgb)

    # 제목
    h = doc.add_heading(data.get("title",""), 1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h.runs:
        r.font.color.rgb = RGBColor(*pr)

    if data.get("subtitle"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(p, data["subtitle"], sz=11)
    doc.add_paragraph()

    # 요약
    if data.get("summary"):
        ar(doc.add_paragraph(), "[ SUMMARY ]", bold=True, rgb=pr)
        ar(doc.add_paragraph(), data["summary"], sz=10)
        doc.add_paragraph()

    # 목차
    ar(doc.add_paragraph(), "[ CONTENTS ]", bold=True, rgb=ac)
    for i, kp in enumerate(data.get("key_points",[]), 1):
        p = doc.add_paragraph()
        ar(p, f"  {i:02d}. {kp.get('heading','')}", bold=True, sz=10, rgb=ac)
        ct = kp.get("content","")
        preview = ct[:80] + "..." if len(ct) > 80 else ct
        ar(doc.add_paragraph(), f"      {preview}", sz=8.5)
    doc.add_paragraph()

    # 상세 본문
    ar(doc.add_paragraph(), "[ DETAILED CONTENT ]", bold=True, rgb=pr)
    doc.add_paragraph()
    for i, kp in enumerate(data.get("key_points",[]), 1):
        ar(doc.add_paragraph(), f"{i:02d}. {kp.get('heading','')}", bold=True, sz=11, rgb=ac)
        ar(doc.add_paragraph(), f"    {kp.get('content','')}", sz=9)
        doc.add_paragraph()

    # 통계
    if data.get("stats"):
        ar(doc.add_paragraph(), "[ 핵심 수치 ]", bold=True, rgb=pr)
        for s in data["stats"]:
            ar(doc.add_paragraph(), f"  {s.get('value','')}  —  {s.get('label','')}", sz=10)
        doc.add_paragraph()

    # 결론
    if data.get("takeaway"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(p, f"결론: {data['takeaway']}", bold=True, sz=10, rgb=pr)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
